from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("deliverables/RAB_Pengembangan_Aplikasi_Menengs_Rinci.docx")
LOGO = Path("public/menengs-logo.png")
NAVY, BLUE, CYAN, LIGHT, INK, MUTED, WHITE = "073B5A", "087FA9", "EAF8FB", "F4F8FB", "173042", "5F7484", "FFFFFF"

def font(run, size=11, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size, run.bold = Pt(size), bold
    run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, color):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)

def margins(cell, top=100, start=120, bottom=100, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if tcMar.getparent() is None: tcPr.append(tcMar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        n = tcMar.find(qn(f"w:{name}")) or OxmlElement(f"w:{name}")
        if n.getparent() is None: tcMar.append(n)
        n.set(qn("w:w"), str(val)); n.set(qn("w:type"), "dxa")

def cell(cell, value, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""; p = cell.paragraphs[0]; p.alignment = align; p.paragraph_format.space_after = Pt(0)
    font(p.add_run(str(value)), size, bold, color); margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def fixed(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW") or OxmlElement("w:tblW")
    if tblW.getparent() is None: tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    layout = tblPr.first_child_found_in("w:tblLayout") or OxmlElement("w:tblLayout")
    if layout.getparent() is None: tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for i, width in enumerate(widths):
        table._tbl.tblGrid.gridCol_lst[i].set(qn("w:w"), str(width))
        for row in table.rows:
            tcW = row.cells[i]._tc.get_or_add_tcPr().first_child_found_in("w:tcW") or OxmlElement("w:tcW")
            if tcW.getparent() is None: row.cells[i]._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa")

def para(doc, text="", size=10.5, bold=False, color=INK, before=0, after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(); p.alignment = align; p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.1
    font(p.add_run(text), size, bold, color); return p

def heading(doc, text, level=1):
    return para(doc, text, 16 if level == 1 else 12.5, True, BLUE if level == 1 else NAVY, 15 if level == 1 else 10, 6)

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.1; font(p.add_run(text), 10.5)

def table(doc, headers, rows, widths, total=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; fixed(t, widths)
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], NAVY); cell(t.rows[0].cells[i], h, True, WHITE, 9.3, WD_ALIGN_PARAGRAPH.CENTER)
    for index, values in enumerate(rows):
        cells = t.add_row().cells
        if index % 2: [shade(c, LIGHT) for c in cells]
        for i, value in enumerate(values):
            cell(cells[i], value, i in (0, len(values) - 1) and i != 3, INK if i != len(values)-1 else INK, 9.3, WD_ALIGN_PARAGRAPH.RIGHT if i == len(values)-1 else (WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT))
    if total:
        cells = t.add_row().cells
        for c in cells: shade(c, CYAN)
        for i, value in enumerate(total):
            cell(cells[i], value, True, NAVY, 10, WD_ALIGN_PARAGRAPH.RIGHT if i == len(total)-2 else WD_ALIGN_PARAGRAPH.LEFT)
    return t

def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document(); sec = doc.sections[0]
    sec.top_margin, sec.bottom_margin, sec.left_margin, sec.right_margin = Inches(.72), Inches(.72), Inches(.78), Inches(.78)
    sec.header_distance, sec.footer_distance = Inches(.32), Inches(.35)
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("RAB Pengembangan Sistem Operasional Menengs  |  VeinTech  |  31 Juli 2026"), 8.5, False, MUTED)

    cover = doc.add_table(rows=1, cols=2); fixed(cover, [2900, 6460]); left, right = cover.rows[0].cells
    shade(left, NAVY); margins(left, 220, 220, 220, 220); margins(right, 250, 260, 250, 260)
    if LOGO.exists():
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; left.paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(1.42))
    else: cell(left, "MENENGS", True, WHITE, 20, WD_ALIGN_PARAGRAPH.CENTER)
    p = right.paragraphs[0]; p.paragraph_format.space_after = Pt(5); font(p.add_run("RAB / PENAWARAN ANGGARAN"), 11, True, BLUE)
    p = right.add_paragraph(); p.paragraph_format.space_after = Pt(7); font(p.add_run("Pengembangan Sistem Operasional Menengs"), 23, True, NAVY)
    p = right.add_paragraph(); p.paragraph_format.space_after = Pt(12); font(p.add_run("Sistem stok, POS outlet, penjualan multi-kanal, dan kendali operasional usaha snack."), 11, False, MUTED)
    p = right.add_paragraph(); p.paragraph_format.space_after = Pt(0); font(p.add_run("Disusun untuk: Menengs  |  Disusun oleh: VeinTech  |  31 Juli 2026"), 9.5, False, INK)
    para(doc, "", after=3)

    heading(doc, "Ringkasan Penawaran")
    para(doc, "Skema yang direkomendasikan adalah kontrak harga tetap berbasis tahapan (fixed-price milestone). Skema ini memberi kepastian anggaran bagi Menengs dan memastikan pembayaran selalu mengikuti hasil kerja yang dapat diuji.")
    summary = doc.add_table(rows=1, cols=3); fixed(summary, [3120, 3120, 3120])
    for c, label, value in zip(summary.rows[0].cells, ["Nilai penawaran", "Estimasi pelaksanaan", "Garansi perbaikan"], ["Rp 30.800.000", "2 bulan kerja", "3 bulan pasca go-live"]):
        shade(c, CYAN); c.text = ""; margins(c, 150, 150, 150, 150)
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(3); font(p.add_run(label.upper()), 8.5, True, MUTED)
        p = c.add_paragraph(); p.paragraph_format.space_after = Pt(0); font(p.add_run(value), 14.5, True, NAVY)
    para(doc, "BONUS INFRASTRUKTUR: Server aplikasi, hosting, dan domain sudah termasuk untuk 12 bulan pertama. Menengs dapat langsung menggunakan aplikasi setelah go-live tanpa perlu menyiapkan layanan tersebut secara terpisah.", 10, True, NAVY, 7, 8)
    heading(doc, "Ruang Lingkup Utama", 2)
    for item in [
        "Master data produk, varian, lokasi, pengguna, peran dan akses berbasis RBAC.",
        "Inventaris real-time: stok masuk, saldo per lokasi, transfer, penerimaan, stock opname, retur, dan riwayat pergerakan.",
        "POS outlet dengan pemilihan varian, scanner barcode, pembayaran tunai/QRIS/transfer, serta kanal online/reseller.",
        "Dashboard, analitik per kanal, laporan operasional, kalkulator HPP, dan biaya marketplace.",
        "API, basis data, pengujian UAT, deployment, pelatihan admin, dan dokumentasi serah-terima."]:
        bullet(doc, item)

    heading(doc, "Fitur yang Akan Diterima Menengs", 1)
    para(doc, "Bagian ini menjelaskan fungsi aplikasi dengan bahasa sederhana. Tujuannya agar seluruh pihak memahami hasil yang akan dipakai sehari-hari oleh owner, admin, gudang, PIC outlet, kasir, dan tim keuangan.")
    table(doc, ["Area", "Fungsi yang diterima", "Manfaat bagi Menengs"], [
        ("Beranda & ringkasan usaha", "Halaman ringkas untuk melihat penjualan, stok, transfer yang masih berjalan, serta stok yang perlu diperhatikan.", "Owner dapat mengetahui kondisi usaha tanpa membuka banyak laporan."),
        ("Produk, varian & foto", "Input produk, pilihan rasa/ukuran/kemasan, harga jual, harga modal, SKU/barcode, dan foto produk.", "Kasir dan gudang memilih barang yang tepat; kesalahan varian berkurang."),
        ("Lokasi usaha", "Mencatat gudang utama, outlet, booth, atau lokasi lain beserta status aktifnya.", "Stok tidak tercampur antar lokasi dan owner bisa melihat posisi barang."),
        ("Akun & hak akses", "Akun berbeda untuk owner, admin, PIC outlet, kasir, gudang, keuangan, dan karyawan; setiap akun hanya melihat/menjalankan tugas yang relevan.", "Data penting lebih aman dan pekerjaan tiap tim lebih terarah."),
        ("Stok masuk & supplier", "Mencatat barang yang datang, asal supplier, jumlah, harga beli, catatan, serta bukti foto bila diperlukan.", "Riwayat pembelian lebih rapi dan stok bertambah dengan dasar yang jelas."),
        ("Stok per lokasi", "Menampilkan saldo stok aktual per produk dan per lokasi, termasuk batas stok minimum.", "Tim cepat tahu barang mana yang perlu dikirim atau dibeli ulang."),
        ("Transfer antar lokasi", "Membuat satu dokumen transfer berisi banyak barang, kode transfer, status perjalanan, bukti, dan konfirmasi penerimaan oleh tujuan.", "Pengiriman gudang ke outlet dapat ditelusuri dan stok tujuan hanya bertambah setelah diterima."),
        ("Stock opname & retur", "Mencatat hasil hitung fisik, selisih, alasan penyesuaian, serta barang yang dikembalikan.", "Selisih stok tidak hilang begitu saja dan dapat ditinjau kembali."),
        ("Riwayat pergerakan stok", "Riwayat barang masuk, keluar, terjual, ditransfer, atau disesuaikan beserta waktu dan referensinya.", "Memudahkan pengecekan saat ada pertanyaan mengenai perubahan stok."),
        ("Kasir / penjualan outlet", "Layar pencatatan penjualan cepat: cari produk atau scan barcode, pilih jumlah, pembayaran tunai/QRIS/transfer, lalu simpan transaksi.", "Melayani antrean lebih cepat sambil stok outlet otomatis berkurang."),
        ("Penjualan online & reseller", "Mencatat transaksi dari kanal online dan reseller dengan kanal pembayaran yang jelas.", "Semua pemasukan tercatat pada satu sistem dan mudah dibandingkan dengan penjualan outlet."),
        ("Dashboard, analitik & laporan", "Grafik dan laporan penjualan per kanal, omzet, produk, stok, serta aktivitas operasional. Laporan dapat diunduh untuk kebutuhan evaluasi.", "Owner memperoleh bahan keputusan yang mudah dibaca, bukan sekadar data mentah."),
        ("Karyawan, absensi & kasbon", "Data karyawan, lokasi penugasan, kehadiran, catatan kasbon/cicilan, dan pencatatan pembayaran gaji beserta bukti transfer.", "Administrasi tim lebih tertib dan pembayaran dapat ditelusuri."),
        ("Kalkulator HPP & biaya marketplace", "Menghitung bahan baku, biaya kemasan/tenaga kerja, hasil produksi, rekomendasi harga jual, diskon, komisi, dan biaya marketplace.", "Menengs dapat menetapkan harga yang tetap menguntungkan setelah biaya platform."),
        ("Akses melalui perangkat", "Tampilan menyesuaikan komputer, tablet, dan ponsel; fitur operasional utama dirancang untuk dipakai di lapangan.", "Tim outlet tidak harus selalu bergantung pada komputer kantor."),
    ], [1750, 4100, 3510])

    heading(doc, "Alur Kerja Sederhana", 2)
    para(doc, "Secara singkat, barang diterima di gudang atau outlet, stoknya tercatat, lalu barang dapat dijual atau dipindahkan ke lokasi lain. Setiap penjualan mengurangi stok lokasi terkait. Jika barang dikirim antar lokasi, stok tujuan baru bertambah setelah PIC tujuan mengonfirmasi penerimaan. Owner melihat seluruh ringkasan melalui dashboard dan laporan.")

    heading(doc, "Bonus Implementasi (Opsional Saat Client Siap)")
    para(doc, "Sebagai nilai tambah, VeinTech menyediakan sesi pendampingan konfigurasi awal BigSeller tanpa biaya tambahan apabila Menengs memutuskan menggunakannya. Cakupan bonus ini meliputi:", 10)
    for item in [
        "Membantu menghubungkan akun marketplace yang didukung BigSeller ke satu dashboard operasional.",
        "Membantu pengaturan dasar sinkronisasi pesanan masuk dan pemetaan stok agar pembaruan stok dapat dijalankan melalui BigSeller.",
        "Pengujian alur pesanan dan stok bersama PIC client pada kanal yang telah terhubung."]:
        bullet(doc, item)
    para(doc, "Catatan: bonus ini tidak menjadi ketergantungan aplikasi Menengs dan tidak termasuk integrasi otomatis marketplace pada tahap pengembangan saat ini. Lisensi/langganan BigSeller, akun marketplace aktif, kredensial integrasi, serta kebijakan dan batasan masing-masing marketplace disediakan oleh client.", 9.5)

    heading(doc, "Bonus Infrastruktur", 2)
    para(doc, "Biaya berikut diberikan sebagai bonus dan sudah termasuk dalam total nilai penawaran Rp 30.800.000. Tidak ada biaya terpisah pada 12 bulan pertama sejak aplikasi go-live.", 10)
    table(doc, ["Bonus", "Cakupan", "Masa berlaku"], [
        ("Server aplikasi", "Penyediaan lingkungan server untuk menjalankan aplikasi Menengs.", "12 bulan pertama"),
        ("Hosting", "Penyimpanan dan akses aplikasi melalui internet untuk kebutuhan operasional Menengs.", "12 bulan pertama"),
        ("Domain", "Satu alamat website/domain untuk akses aplikasi Menengs.", "12 bulan pertama"),
    ], [2300, 4900, 2160])
    para(doc, "Setelah masa bonus berakhir, biaya perpanjangan server, hosting, dan domain mengikuti harga penyedia layanan pada saat perpanjangan. Biaya tersebut akan diinformasikan terlebih dahulu kepada Menengs sebelum jatuh tempo.", 9.5, False, MUTED)

    heading(doc, "Rincian Anggaran Biaya (RAB)")
    para(doc, "Nilai mencakup analisis, pengembangan, pengujian, dan serah-terima versi 1.0. Sebagai dukungan implementasi Menengs, diberikan diskon pada setiap komponen. Harga bersifat fixed-price selama ruang lingkup tidak berubah.", 10)
    table(doc, ["No.", "Komponen pekerjaan", "Nilai awal", "Diskon", "Nilai penawaran"], [
        ("1", "Analisis kebutuhan, perancangan alur operasional & UI/UX", "Rp 4.500.000", "Rp 1.800.000", "Rp 2.700.000"),
        ("2", "Master data, pengguna, dan RBAC", "Rp 7.500.000", "Rp 3.000.000", "Rp 4.500.000"),
        ("3", "Modul inventaris dan sinkronisasi antar lokasi", "Rp 9.500.000", "Rp 3.700.000", "Rp 5.800.000"),
        ("4", "POS outlet & penjualan multi-kanal", "Rp 8.000.000", "Rp 3.200.000", "Rp 4.800.000"),
        ("5", "Dashboard, analitik, laporan, HPP & marketplace", "Rp 6.500.000", "Rp 2.600.000", "Rp 3.900.000"),
        ("6", "API, database, keamanan, dan integrasi internal", "Rp 7.500.000", "Rp 3.000.000", "Rp 4.500.000"),
        ("7", "QA, UAT, deployment, dan pelatihan", "Rp 5.000.000", "Rp 2.000.000", "Rp 3.000.000"),
        ("8", "Dokumentasi dan serah-terima teknis", "Rp 2.300.000", "Rp 700.000", "Rp 1.600.000"),
    ], [500, 3850, 1650, 1550, 1810], ("", "TOTAL NILAI PENGEMBANGAN", "Rp 50.800.000", "Rp 20.000.000", "Rp 30.800.000"))

    heading(doc, "Tahapan Pembayaran")
    para(doc, "Pembayaran berbasis hasil kerja agar kedua pihak memiliki kontrol jelas atas progres proyek.", 10)
    table(doc, ["Tahap", "Output / dasar penagihan", "Porsi", "Nilai"], [
        ("I", "Persetujuan RAB, SOW, dan kickoff", "30%", "Rp 9.240.000"),
        ("II", "Master data, RBAC, inventaris, dan POS siap UAT", "35%", "Rp 10.780.000"),
        ("III", "Analitik, HPP, marketplace, dan UAT selesai", "25%", "Rp 7.700.000"),
        ("IV", "Go-live, pelatihan, dan serah-terima", "10%", "Rp 3.080.000"),
    ], [1000, 3600, 1700, 3060])

    heading(doc, "Jadwal Pelaksanaan")
    table(doc, ["Periode", "Fokus pekerjaan", "Output utama"], [
        ("Minggu 1–2", "Discovery, data model, UI/UX, dan master data", "Prototype disetujui"),
        ("Minggu 3–4", "Inventaris, transfer, stok masuk, RBAC", "Alur stok siap UAT"),
        ("Minggu 5–6", "POS, penjualan, dashboard, dan analitik", "Modul transaksi siap UAT"),
        ("Minggu 7–8", "HPP, marketplace, QA, deployment, training", "Go-live & handover"),
    ], [1650, 5550, 2160])

    heading(doc, "Garansi, Dukungan, dan Opsi Pemeliharaan")
    para(doc, "Garansi perbaikan bug selama 3 bulan setelah go-live sudah termasuk. Perubahan proses bisnis baru ditangani melalui change request.", 10)
    table(doc, ["Pilihan", "Cakupan", "Nilai"], [
        ("Garansi pasca go-live", "Perbaikan bug dalam ruang lingkup selama 3 bulan", "Termasuk"),
        ("Pemeliharaan bulanan (opsional)", "Monitoring, backup review, dan perbaikan minor maksimal 8 jam/bulan", "Rp 1.850.000/bulan"),
    ], [3150, 3700, 2510])

    heading(doc, "Asumsi dan Pengecualian")
    for item in [
        "Nilai belum termasuk PPN, SMS/WhatsApp, biaya payment gateway, atau layanan pihak ketiga lain di luar bonus infrastruktur 12 bulan pertama.",
        "Integrasi otomatis marketplace, kurir, dan payment gateway tidak termasuk kecuali disepakati melalui change request.",
        "Client menyediakan data awal produk, lokasi, pengguna, kebijakan operasional, dan PIC UAT sesuai jadwal.",
        "Perubahan kebutuhan di luar ruang lingkup dihitung melalui change request tertulis.",
        "Penawaran berlaku 14 hari kalender sejak tanggal dokumen ini."]:
        bullet(doc, item)

    heading(doc, "Persetujuan")
    para(doc, "Dengan menandatangani bagian berikut, para pihak menyetujui ruang lingkup, nilai, dan skema pembayaran dalam RAB ini.", 10)
    sign = doc.add_table(rows=1, cols=2); fixed(sign, [4680, 4680])
    for c, title in zip(sign.rows[0].cells, ["Pihak Menengs", "Penyedia / VeinTech"]):
        c.text = ""; margins(c, 150, 150, 150, 150)
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        font(p.add_run(title + "\n(______________________________)"), 11, True, NAVY)
    doc.save(OUT); print(OUT)

if __name__ == "__main__":
    build()
