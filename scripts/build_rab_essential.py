from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("deliverables/RAB_Pengembangan_Aplikasi_Menengs_Essential_Rp10Juta.docx")

NAVY = "073B5A"
BLUE = "087FA9"
CYAN = "DDF6FB"
LIGHT = "F3F8FA"
INK = "17364A"
MUTED = "5A7485"
LINE = "C9E0E8"
GREEN = "137A50"
AMBER = "A05A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, value, size=10.5, bold=False, color=INK, after=6, before=0, align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    set_run(p.add_run(value), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, value, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(value), size=15 if level == 1 else 12, bold=True, color=NAVY)
    return p


def table_header(table, labels):
    row = table.rows[0]
    set_repeat_table_header(row)
    for cell, label in zip(row.cells, labels):
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=100, bottom=100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label), size=9, bold=True, color="FFFFFF")


def add_bullet(doc, value):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    set_run(p.add_run(value), size=10.2, color=INK)
    return p


def add_callout(doc, title, detail, fill=CYAN, title_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    set_cell_border(cell, top={"val": "single", "sz": 8, "color": LINE}, bottom={"val": "single", "sz": 8, "color": LINE}, left={"val": "single", "sz": 8, "color": LINE}, right={"val": "single", "sz": 8, "color": LINE})
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title), size=11, bold=True, color=title_color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(detail), size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("RAB Menengs Essential • VeinTech • 1 Agustus 2026"), size=8.5, color=MUTED)


def write_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.68)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.32)
    footer(sec)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)

    # Cover / executive offer
    add_text(doc, "VEINTECH", size=11, bold=True, color=BLUE, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "RAB PENGEMBANGAN APLIKASI", size=24, bold=True, color=NAVY, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "MENENGS — PAKET ESSENTIAL", size=15, bold=True, color=BLUE, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Sistem dasar untuk mencatat stok antar lokasi dan penjualan outlet secara rapi.", size=11, color=MUTED, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)

    meta = doc.add_table(rows=2, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(meta, [3120, 3120, 3120])
    meta_data = [
        ("NILAI PENGEMBANGAN", "Rp 10.000.000"),
        ("WAKTU PENGERJAAN", "1 bulan kerja"),
        ("INFRASTRUKTUR", "Server, hosting & domain 1 tahun"),
    ]
    for idx, (label, value) in enumerate(meta_data):
        c1, c2 = meta.cell(0, idx), meta.cell(1, idx)
        for c in (c1, c2):
            set_cell_margins(c, top=110, bottom=110)
            set_cell_border(c, top={"val": "single", "sz": 8, "color": LINE}, bottom={"val": "single", "sz": 8, "color": LINE}, left={"val": "single", "sz": 8, "color": LINE}, right={"val": "single", "sz": 8, "color": LINE})
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(c1, LIGHT)
        c1.text = ""
        p = c1.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label), size=8.5, bold=True, color=MUTED)
        c2.text = ""
        p = c2.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), size=11.5, bold=True, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_heading(doc, "Ringkasan penawaran")
    add_text(doc, "Paket Essential dibuat untuk kebutuhan operasional inti Menengs dengan anggaran yang terukur. Fokusnya adalah data produk, stok per lokasi, transfer stok, dan kasir/penjualan outlet. Seluruh fitur dirancang agar mudah digunakan oleh owner maupun tim outlet.", after=7)
    add_callout(doc, "Yang termasuk dalam nilai Rp10.000.000", "Analisis kebutuhan dasar, pengembangan aplikasi, pengujian bersama (UAT), publikasi aplikasi, pelatihan singkat admin, serta server, hosting, dan satu domain selama 12 bulan pertama.")

    add_heading(doc, "Hasil yang diterima")
    for item in [
        "Aplikasi web yang dapat dibuka dari komputer, tablet, dan ponsel.",
        "Akun owner dan akun tim dengan akses sesuai peran kerja.",
        "Pencatatan stok dan penjualan yang tersimpan di basis data aplikasi.",
        "Serah-terima aplikasi setelah UAT disetujui oleh PIC client.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "Ruang lingkup fitur dasar")
    add_text(doc, "Berikut adalah fitur yang dikerjakan pada paket ini. Penjelasan dibuat sederhana agar mudah dipakai sebagai acuan bersama saat UAT.", after=7)
    feature_rows = [
        ("Produk & varian", "Menambah produk, rasa/ukuran/kemasan, harga jual, SKU/barcode, dan status aktif produk.", "Barang lebih mudah ditemukan dan risiko salah pilih varian berkurang."),
        ("Lokasi & akun tim", "Mencatat gudang/outlet serta membuat akun owner dan tim dengan akses sesuai tugas.", "Data operasional lebih rapi dan akses penting tidak terbuka untuk semua orang."),
        ("Stok masuk", "Mencatat barang yang masuk ke gudang atau outlet beserta jumlah dan catatan dasar.", "Saldo stok bertambah dengan riwayat yang jelas."),
        ("Stok per lokasi", "Melihat saldo stok per produk pada masing-masing lokasi serta batas stok minimum.", "Owner dan PIC cepat mengetahui barang yang perlu disiapkan."),
        ("Transfer stok", "Membuat dokumen pengiriman dari satu lokasi ke lokasi lain dan mengonfirmasi penerimaan.", "Stok tujuan bertambah hanya setelah barang diterima."),
        ("POS outlet", "Mencatat penjualan offline, memilih produk/varian, jumlah, dan metode bayar tunai/QRIS/transfer.", "Kasir dapat melayani transaksi harian dan stok outlet otomatis berkurang."),
        ("Dashboard & laporan dasar", "Ringkasan penjualan, stok, dan transaksi utama untuk periode tanggal yang dipilih.", "Owner memiliki gambaran usaha tanpa membuka catatan manual."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(t, [1800, 4200, 3360])
    table_header(t, ["FITUR", "YANG DAPAT DILAKUKAN", "MANFAAT"])
    for idx, row_data in enumerate(feature_rows):
        cells = t.add_row().cells
        for cell, value in zip(cells, row_data):
            set_cell_shading(cell, "FFFFFF" if idx % 2 == 0 else LIGHT)
            set_cell_margins(cell)
            set_cell_border(cell, top={"val": "single", "sz": 5, "color": LINE}, bottom={"val": "single", "sz": 5, "color": LINE}, left={"val": "single", "sz": 5, "color": LINE}, right={"val": "single", "sz": 5, "color": LINE})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(value), size=9.1, bold=(cell == cells[0]), color=INK)

    add_heading(doc, "Batasan paket Essential")
    add_text(doc, "Agar waktu satu bulan dan anggaran Rp10 juta dapat dipenuhi, fitur di bawah ini tidak termasuk dalam penawaran ini:", after=4)
    for item in [
        "Integrasi BigSeller, marketplace, kurir, payment gateway, atau layanan pihak ketiga lainnya.",
        "Pencatatan pesanan marketplace secara otomatis, sinkronisasi stok marketplace, dan biaya marketplace.",
        "Kalkulator HPP/produksi, analitik lanjutan, ekspor laporan khusus, HR/karyawan, absensi, penggajian, dan kasbon.",
        "Pembuatan fitur baru, perubahan alur bisnis, serta revisi setelah UAT dan serah-terima disetujui.",
    ]:
        add_bullet(doc, item)

    add_callout(doc, "Catatan infrastruktur", "Server aplikasi, hosting, dan satu domain disediakan untuk 12 bulan pertama sejak aplikasi go-live. Perpanjangan setelah periode tersebut akan mengikuti biaya penyedia layanan pada saat perpanjangan.", fill="EEF8F2", title_color=GREEN)

    doc.add_page_break()

    add_heading(doc, "Rincian anggaran biaya")
    add_text(doc, "Harga bersifat fixed-price untuk ruang lingkup paket Essential pada dokumen ini.", after=7)
    budget_rows = [
        ("1", "Analisis kebutuhan dasar & rancangan tampilan", "Rp 1.400.000"),
        ("2", "Master produk, lokasi, akun tim & hak akses", "Rp 1.300.000"),
        ("3", "Stok masuk, stok per lokasi & transfer stok", "Rp 2.000.000"),
        ("4", "POS/penjualan offline dan pencatatan pembayaran", "Rp 1.800.000"),
        ("5", "Dashboard serta laporan operasional dasar", "Rp 1.000.000"),
        ("6", "Pengujian UAT, deployment & pelatihan admin", "Rp 1.000.000"),
        ("7", "Server, hosting & domain selama 1 tahun", "Rp 1.500.000"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(t, [620, 6070, 2670])
    table_header(t, ["NO", "KOMPONEN PEKERJAAN", "NILAI"])
    for idx, row_data in enumerate(budget_rows):
        cells = t.add_row().cells
        for col, (cell, value) in enumerate(zip(cells, row_data)):
            set_cell_shading(cell, "FFFFFF" if idx % 2 == 0 else LIGHT)
            set_cell_margins(cell)
            set_cell_border(cell, top={"val": "single", "sz": 5, "color": LINE}, bottom={"val": "single", "sz": 5, "color": LINE}, left={"val": "single", "sz": 5, "color": LINE}, right={"val": "single", "sz": 5, "color": LINE})
            cell.text = ""
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(value), size=9.5, bold=(col == 2), color=INK)
    cells = t.add_row().cells
    for cell in cells:
        set_cell_shading(cell, CYAN)
        set_cell_margins(cell, top=120, bottom=120)
        set_cell_border(cell, top={"val": "single", "sz": 8, "color": BLUE}, bottom={"val": "single", "sz": 8, "color": BLUE}, left={"val": "single", "sz": 8, "color": BLUE}, right={"val": "single", "sz": 8, "color": BLUE})
        cell.text = ""
    p = cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_run(p.add_run(""), size=9)
    p = cells[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_run(p.add_run("TOTAL NILAI PENGEMBANGAN"), size=10.5, bold=True, color=NAVY)
    p = cells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_run(p.add_run("Rp 10.000.000"), size=11.5, bold=True, color=NAVY)

    add_heading(doc, "Jadwal pelaksanaan — 1 bulan")
    schedule = [
        ("Minggu 1", "Kickoff, kebutuhan dasar, rancangan layar, dan konfigurasi master data."),
        ("Minggu 2", "Pengerjaan stok masuk, stok per lokasi, dan transfer antar lokasi."),
        ("Minggu 3", "Pengerjaan POS outlet, dashboard, serta laporan dasar."),
        ("Minggu 4", "UAT bersama PIC, perbaikan hasil UAT yang disetujui sebelum go-live, deployment, dan pelatihan singkat."),
    ]
    for week, detail in schedule:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(f"{week} — "), size=10.2, bold=True, color=BLUE)
        set_run(p.add_run(detail), size=10.2, color=INK)

    add_heading(doc, "Tahapan pembayaran")
    payments = doc.add_table(rows=1, cols=3)
    payments.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(payments, [1300, 5150, 2910])
    table_header(payments, ["TAHAP", "DASAR PENAGIHAN", "NILAI"])
    for idx, row_data in enumerate([
        ("I · 50%", "Persetujuan RAB, ruang lingkup, dan kickoff proyek.", "Rp 5.000.000"),
        ("II · 30%", "Modul stok dan POS dasar siap diuji oleh PIC client.", "Rp 3.000.000"),
        ("III · 20%", "UAT disetujui, aplikasi go-live, dan serah-terima.", "Rp 2.000.000"),
    ]):
        cells = payments.add_row().cells
        for col, (cell, value) in enumerate(zip(cells, row_data)):
            set_cell_shading(cell, "FFFFFF" if idx % 2 == 0 else LIGHT)
            set_cell_margins(cell)
            set_cell_border(cell, top={"val": "single", "sz": 5, "color": LINE}, bottom={"val": "single", "sz": 5, "color": LINE}, left={"val": "single", "sz": 5, "color": LINE}, right={"val": "single", "sz": 5, "color": LINE})
            cell.text = ""
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(value), size=9.2, bold=(col != 1), color=INK)

    # Ketentuan dan persetujuan sengaja ditempatkan bersama agar halaman
    # tanda tangan tidak terpisah sebagai halaman kosong.
    doc.add_page_break()
    add_heading(doc, "Ketentuan penting")
    for item in [
        "Tidak ada masa garansi perbaikan, revisi gratis, maupun dukungan bulanan setelah serah-terima/UAT disetujui.",
        "Revisi yang masih berada dalam daftar UAT dapat dilakukan selama periode UAT sebelum persetujuan serah-terima. Perubahan setelahnya dihitung sebagai pekerjaan tambahan berdasarkan persetujuan tertulis.",
        "Nilai belum termasuk PPN (jika berlaku), perangkat kasir, biaya QRIS/payment gateway, layanan pihak ketiga, serta biaya perpanjangan infrastruktur setelah 12 bulan.",
        "Client menyediakan data produk, lokasi, daftar pengguna, dan PIC pengujian sesuai jadwal agar target satu bulan dapat tercapai.",
        "Penawaran berlaku 14 hari kalender sejak tanggal dokumen.",
    ]:
        add_bullet(doc, item)

    add_text(doc, "Persetujuan", size=12, bold=True, color=NAVY, before=16, after=4)
    add_text(doc, "Dengan menandatangani dokumen ini, para pihak menyetujui ruang lingkup, nilai, jadwal, dan ketentuan pada RAB Paket Essential.", size=10.2, color=MUTED, after=12)
    signatures = doc.add_table(rows=2, cols=2)
    signatures.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(signatures, [4680, 4680])
    for cell, label in zip(signatures.rows[0].cells, ["Pihak Menengs", "Penyedia — VeinTech"]):
        set_cell_margins(cell, top=100, bottom=100)
        cell.text = ""
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label), size=10.5, bold=True, color=INK)
    for cell in signatures.rows[1].cells:
        set_cell_margins(cell, top=450, bottom=120)
        cell.text = ""
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run("Nama / tanda tangan / tanggal"), size=9.5, color=MUTED, italic=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    write_doc()
