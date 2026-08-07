from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "Panduan_UAT_Tester_Menengs_2026-08-04.docx"

NAVY = "0B2545"
BLUE = "0B87A8"
HEAD_BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_CYAN = "EAF7FA"
LIGHT = "F4F6F9"
GRAY = "5F6B7A"
GRID = "C9D6E2"
GREEN = "1D6F42"
AMBER = "7A5A00"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def font(run, size=11, bold=False, color=NAVY, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_p(doc, text="", bold_prefix=None, color=NAVY, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        font(p.add_run(bold_prefix), bold=True, color=color)
        font(p.add_run(text[len(bold_prefix):]), color=color)
    else:
        font(p.add_run(text), color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text))
    return p


def add_callout(doc, title, body, tone="info"):
    fills = {"info": PALE_CYAN, "warn": "FFF4D6", "risk": "FCE8E8", "ok": "EAF6EF"}
    colors = {"info": BLUE, "warn": AMBER, "risk": RED, "ok": GREEN}
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills[tone])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(title), bold=True, color=colors[tone])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    font(p2.add_run(body), size=10.5, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, font_size=9, header_fill=PALE_BLUE, repeat=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed(table, widths)
    hdr = table.rows[0]
    for i, value in enumerate(headers):
        set_cell_shading(hdr.cells[i], header_fill)
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(str(value)), size=font_size, bold=True, color=DARK_BLUE)
    if repeat:
        tr_pr = hdr._tr.get_or_add_trPr()
        rep = OxmlElement("w:tblHeader")
        rep.set(qn("w:val"), "true")
        tr_pr.append(rep)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            font(p.add_run(str(value)), size=font_size, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def page_break(doc):
    doc.add_page_break()


def set_repeat_numbering_styles(doc):
    for style_name, left, hanging in (("List Bullet", 270, 270), ("List Number", 270, 270), ("List Bullet 2", 540, 270)):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        pf = style.paragraph_format
        pf.left_indent = Pt(left / 20)
        pf.first_line_indent = Pt(-hanging / 20)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.25


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, HEAD_BLUE, 18, 10),
        ("Heading 2", 13, HEAD_BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    set_repeat_numbering_styles(doc)

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font(hp.add_run("MENENGS  |  PAKET UAT TESTER"), size=9, bold=True, color=GRAY)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(fp.add_run("Internal • 4 Agustus 2026  |  "), size=8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)

    # Customer-pack opening
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run("USER ACCEPTANCE TESTING"), size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    font(p.add_run("Panduan UAT & Serah Terima Aplikasi Menengs"), size=27, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    font(p.add_run("Scope, acceptance criteria, business flow, test data, API, role-permission, dan batasan sistem."), size=13, color=GRAY)
    add_table(doc, ["Disiapkan untuk", "Lingkungan", "Versi dokumen", "Tanggal"], [["Tim Tester Menengs", "Production terisolasi", "1.0", "4 Agustus 2026"]], [2340, 2340, 2340, 2340], 9.5, LIGHT)
    add_callout(doc, "Tujuan dokumen", "Menjadi satu sumber acuan tester untuk mengeksekusi UAT dari kondisi data kosong, mencatat bukti, menguji pembatasan role, dan memberikan rekomendasi Go / No-Go sebelum serah terima.", "info")
    doc.add_heading("Cara menggunakan dokumen ini", level=2)
    add_number(doc, "Gunakan hanya akun Tester untuk pengujian rutin dan data sintetis.")
    add_number(doc, "Jalankan skenario sesuai urutan business flow agar saldo stok dapat ditelusuri.")
    add_number(doc, "Isi kolom hasil dengan Lulus, Gagal, atau Blocked; lampirkan screenshot dan waktu kejadian.")
    add_number(doc, "Buat issue terpisah untuk setiap kegagalan dan cantumkan ID acceptance criterion.")
    add_number(doc, "Rekomendasikan Go hanya jika seluruh kriteria Critical dan High lulus tanpa defect terbuka.")

    doc.add_heading("Daftar isi", level=2)
    toc = [
        "1. Scope dan Acceptance Criteria",
        "2. User Flow / Business Flow",
        "3. Daftar Akun dan Test Data",
        "4. Dokumentasi API",
        "5. Matriks Role dan Permission",
        "6. Known Issues dan Limitation",
        "7. Form Ringkasan Hasil dan Sign-off",
    ]
    for item in toc:
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("1. Scope dan Acceptance Criteria", level=1)
    add_p(doc, "Lingkup UAT mencakup alur operasional end-to-end, integritas saldo stok, audit trail, role-based access control, pengalaman desktop/mobile, dan kesiapan operasional production. Integrasi eksternal berbayar atau sinkronisasi marketplace otomatis tidak termasuk.")
    doc.add_heading("1.1 Ruang lingkup fungsi", level=2)
    scopes = [
        ("Akses & organisasi", "Login, sesi, lupa/reset password, profil usaha, isolasi workspace"),
        ("Master data", "Lokasi, produk, varian, SKU/barcode, supplier, minimum stok"),
        ("Persediaan", "Stok masuk, saldo per lokasi, transfer, opname, histori, retur"),
        ("Transaksi", "Penjualan offline/online, pembatalan, laporan dan ekspor"),
        ("Pengiriman", "Scan/manual resi, Siap Diangkut, Serah Terima, riwayat"),
        ("HPP & marketplace", "Resep, biaya tambahan, penerapan HPP ke varian, simulasi fee"),
        ("Tim & HR", "Akun staf, role/lokasi, karyawan, absensi, kasbon, payroll"),
        ("Nonfungsional", "Responsive mobile, validasi input, auditability, error handling"),
    ]
    add_table(doc, ["Area", "Cakupan"], scopes, [2700, 6660], 9.5)
    doc.add_heading("1.2 Kriteria kelulusan keseluruhan", level=2)
    add_bullet(doc, "100% acceptance criterion Critical lulus.")
    add_bullet(doc, "Minimal 95% acceptance criterion High dan Medium lulus; tidak ada defect High/Critical terbuka.")
    add_bullet(doc, "Saldo stok akhir cocok dengan perhitungan manual dan setiap perubahan memiliki histori.")
    add_bullet(doc, "Role terbatas tidak dapat membaca atau mengubah data di luar scope lokasi/izin.")
    add_bullet(doc, "Tidak ada error console/API berulang yang menghambat alur utama.")

    criteria = [
        ("AUTH-01", "Critical", "Login akun aktif berhasil; kredensial salah ditolak tanpa membocorkan akun."),
        ("AUTH-02", "High", "Sesi kedaluwarsa mengarahkan user untuk login kembali; token berlaku 12 jam."),
        ("ORG-01", "Critical", "Akun Tester hanya melihat data workspace Meneng Testing dan mulai dari data kosong."),
        ("MD-01", "High", "Owner dapat membuat lokasi Gudang dan Outlet; duplikat nama+jenis ditolak."),
        ("MD-02", "Critical", "Produk/varian tersimpan dengan unit, SKU/barcode unik, harga, dan status aktif."),
        ("INV-01", "Critical", "Stok masuk menambah saldo lokasi dan membuat histori dengan user/waktu."),
        ("INV-02", "Critical", "Transfer mengurangi asal saat dikirim dan menambah tujuan hanya saat diterima."),
        ("INV-03", "Critical", "Opname menghitung selisih terhadap saldo sistem dan tidak menggandakan koreksi saat diulang/edit."),
        ("INV-04", "Critical", "Pembatalan stok masuk/transfer/opname/retur membalik saldo tepat satu kali dan menyimpan alasan."),
        ("SALE-01", "Critical", "Penjualan mengurangi stok sesuai item dan total nominal tersimpan benar."),
        ("SALE-02", "Critical", "Void penjualan mengembalikan stok dan status transaksi menjadi dibatalkan."),
        ("SHIP-01", "High", "Resi valid dapat dicatat melalui kamera atau manual dan paket berpindah ke Siap Diangkut."),
        ("SHIP-02", "High", "Scan serah-terima dan finalisasi batch mengubah status ke handed-over tanpa duplikasi resi."),
        ("HPP-01", "Critical", "Total bahan, biaya tambahan, HPP total dan HPP per hasil dihitung konsisten termasuk angka desimal."),
        ("HPP-02", "Critical", "Penerapan HPP ke varian yang dipilih memperbarui Harga Modal setelah konfirmasi dampak."),
        ("HPP-03", "High", "Simulasi marketplace menghitung diskon, fee, affiliate, net payout, laba/rugi, dan impas."),
        ("RBAC-01", "Critical", "Owner/Admin/PIC/Gudang/Kasir/Keuangan/Karyawan hanya melihat menu dan aksi yang diizinkan."),
        ("RBAC-02", "Critical", "PIC/Gudang/Kasir tidak dapat mengakses lokasi lain lewat UI maupun request API langsung."),
        ("UX-01", "High", "Semua input nominal menampilkan format Rupiah tanpa mengubah nilai numerik tersimpan."),
        ("UX-02", "High", "Input angka tidak menyimpan nol di depan; input kuantitas yang mendukung pecahan menerima desimal."),
        ("UX-03", "Medium", "Daftar mobile mendekat ke filter saat date range tertutup dan bergeser saat panel dibuka."),
        ("NFR-01", "High", "Tidak ada saldo negatif, double-submit, atau transaksi parsial saat validasi gagal."),
        ("NFR-02", "Medium", "Halaman utama tetap usable pada 360px mobile dan desktop 1366px tanpa overlap."),
    ]
    doc.add_heading("1.3 Checklist acceptance", level=2)
    add_p(doc, "Isi Hasil dengan L = Lulus, G = Gagal, B = Blocked. Bukti minimal berupa screenshot, ID transaksi, dan timestamp.", color=GRAY)
    add_table(doc, ["ID", "Prioritas", "Acceptance criterion", "Hasil", "Bukti / Issue"], criteria, [820, 980, 5200, 760, 1600], 8.3)

    page_break(doc)
    doc.add_heading("2. User Flow / Business Flow", level=1)
    add_callout(doc, "Urutan yang direkomendasikan", "Buat master Produk & Varian terlebih dahulu, lalu susun HPP dan hubungkan ke varian. Setelah itu isi stok awal/stok masuk dan jalankan transaksi harian.", "ok")
    flows = [
        ("A", "Persiapan workspace", ["Login sebagai Tester (Owner)", "Lengkapi profil usaha uji", "Periksa gudang awal", "Buat Outlet Test", "Buat akun role uji bila diperlukan"]),
        ("B", "Produk dan HPP", ["Buat produk + varian", "Tentukan satuan stok", "Buat resep HPP", "Isi bahan dan biaya tambahan", "Pilih varian tujuan", "Konfirmasi sinkronisasi Harga Modal"]),
        ("C", "Siklus persediaan", ["Stok masuk ke Gudang Test", "Periksa saldo & histori", "Kirim transfer ke Outlet Test", "Terima di outlet", "Lakukan opname", "Uji edit/pembatalan"]),
        ("D", "Penjualan dan retur", ["Catat penjualan outlet", "Pastikan stok berkurang", "Catat retur pelanggan/supplier", "Void transaksi uji", "Cocokkan saldo dan laporan"]),
        ("E", "Pengiriman", ["Pilih lokasi + marketplace", "Scan atau input resi", "Verifikasi Siap Diangkut", "Scan resi ke batch serah-terima", "Finalisasi batch", "Verifikasi Riwayat"]),
        ("F", "Tim dan payroll", ["Buat akun staf", "Tetapkan role dan lokasi", "Hubungkan data karyawan", "Uji absensi", "Catat kasbon", "Proses payroll dan verifikasi akses"]),
    ]
    for code, title, steps in flows:
        doc.add_heading(f"{code}. {title}", level=2)
        for index, step in enumerate(steps, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            font(p.add_run(f"{code}{index}  "), bold=True, color=DARK_BLUE)
            font(p.add_run(step), color=NAVY)
    doc.add_heading("2.1 Dampak penerapan HPP ke varian", level=2)
    add_p(doc, "Satu resep HPP dapat diterapkan ke beberapa varian. Saat disimpan dan dikonfirmasi, nilai HPP per hasil menjadi Harga Modal pada semua varian yang dipilih. Harga Jual, Harga Reseller, stok, SKU, dan barcode tidak berubah. Jika resep HPP diubah kemudian, tester wajib memeriksa ulang apakah sinkronisasi ke varian dilakukan sesuai konfirmasi aplikasi.")

    page_break(doc)
    doc.add_heading("3. Daftar Akun dan Test Data", level=1)
    add_callout(doc, "Aturan keamanan", "Dokumen yang dibagikan ke tester hanya memuat password akun Tester. Password Owner Serah Terima dan Developer dikelola oleh pemilik proyek dan tidak boleh dikirim melalui grup umum.", "warn")
    accounts = [
        ("Tester", "tester@menengsofficial.com", "Tst!Meneng#2026_Qa7", "Owner", "Meneng Testing", "UAT dari data kosong"),
        ("Owner serah terima", "owner@menengsofficial.com", "Akses terbatas", "Owner", "Meneng Serah Terima", "Dipakai setelah Go-Live"),
        ("Developer", "developer@menengsofficial.com", "Akses terbatas", "Owner", "Meneng Development", "Debug dan verifikasi teknis"),
    ]
    add_table(doc, ["Akun", "Email", "Password", "Role", "Workspace", "Kegunaan"], accounts, [1050, 2200, 1800, 850, 1500, 1960], 8.4)
    add_p(doc, "URL production: https://internal.menengsofficial.com", bold_prefix="URL production:")
    add_p(doc, "Kondisi awal workspace Tester: 1 gudang bawaan, 0 produk, 0 saldo, 0 penjualan, 0 transfer, 0 movement, 0 opname, 0 stok masuk, 0 retur, dan 0 pengiriman.")

    doc.add_heading("3.1 Dataset UAT yang harus dibuat tester", level=2)
    test_data = [
        ("Lokasi", "Gudang Test; Outlet Test", "Gudang sebagai sumber, outlet sebagai tujuan"),
        ("Produk", "Produk UAT", "Satuan stok Pcs"),
        ("Varian", "Basic; Premium", "SKU UAT-BSC-001 dan UAT-PRM-002"),
        ("Supplier", "Supplier UAT", "Tidak memakai data supplier nyata"),
        ("HPP", "Yield 10 Pcs; bahan Rp50.000; biaya tambahan Rp10.000", "HPP total Rp60.000; HPP/Pcs Rp6.000"),
        ("Stok masuk", "100 Pcs Basic ke Gudang Test", "Saldo Gudang = 100"),
        ("Transfer", "20 Pcs Gudang → Outlet", "Setelah diterima: Gudang 80; Outlet 20"),
        ("Penjualan", "2 Pcs Basic @ Rp10.000", "Outlet menjadi 18; total Rp20.000"),
        ("Retur pelanggan", "1 Pcs Basic", "Outlet menjadi 19"),
        ("Opname", "Fisik 18 Pcs; alasan UAT selisih", "Selisih -1; saldo Outlet menjadi 18"),
        ("Pengiriman", "UAT-RESI-20260804-001", "Marketplace Test; gunakan resi sintetis"),
    ]
    add_table(doc, ["Objek", "Data uji", "Expected"], test_data, [1600, 4200, 3560], 8.8)
    add_callout(doc, "Larangan data produksi", "Jangan memasukkan nama pelanggan, nomor telepon, alamat, resi nyata, bukti transfer, atau foto identitas asli. Gunakan prefiks UAT pada seluruh nama dan catatan.", "risk")

    page_break(doc)
    doc.add_heading("4. Dokumentasi API", level=1)
    add_p(doc, "Base URL production: https://internal.menengsofficial.com/api", bold_prefix="Base URL production:")
    add_p(doc, "Semua endpoint selain login, reset password, dan health memerlukan header Authorization: Bearer <token>. Body menggunakan application/json kecuali upload image. Token login berlaku 12 jam.")
    api_rows = [
        ("POST", "/login", "Publik", "Login; menghasilkan token dan profil user"),
        ("POST", "/auth/forgot-password", "Publik", "Memulai reset password; respons tidak membocorkan email terdaftar"),
        ("POST", "/auth/reset-password", "Publik", "Mengubah password dengan token reset"),
        ("PATCH", "/profile/password", "Login", "Mengubah password sendiri"),
        ("GET", "/health", "Publik", "Health check aplikasi dan database"),
        ("GET", "/state", "Login", "Mengambil state sesuai organisasi, role, dan lokasi"),
        ("PUT", "/organization", "Owner", "Memperbarui profil usaha"),
        ("POST/PATCH", "/users, /users/:id", "Owner", "Membuat/mengubah akun, role, lokasi, status"),
        ("POST/PATCH", "/commands/products, /commands/products/:id", "Owner/Admin", "Membuat/mengubah produk dan varian"),
        ("POST/PATCH", "/commands/locations, /commands/locations/:id", "Owner/Admin", "Membuat/mengubah lokasi"),
        ("POST/PATCH", "/commands/suppliers, /commands/suppliers/:id", "Owner/Admin", "Membuat/mengubah supplier"),
        ("POST", "/commands/receipts", "Stock In", "Mencatat stok masuk multi-item"),
        ("PATCH", "/commands/receipts/:id", "Stock In", "Merevisi satu baris stok masuk"),
        ("POST", "/commands/transfers", "Transfer Send", "Mengirim stok dari lokasi asal"),
        ("POST", "/commands/transfers/:code/receive", "Transfer Receive", "Menerima dokumen transfer"),
        ("POST/PATCH", "/commands/opnames, /commands/opnames/:id", "Stock Opname", "Mencatat/merevisi hasil fisik"),
        ("POST", "/commands/sales", "Sale Create", "Mencatat penjualan dan mengurangi stok"),
        ("POST", "/commands/returns", "Stock Out", "Retur pelanggan atau supplier"),
        ("POST", "/commands/cancel", "Sesuai objek", "Membatalkan sale/transfer/receipt/opname/return"),
        ("POST", "/commands/shipping/ready", "Shipping Manage", "Mencatat resi sebagai siap diangkut"),
        ("POST", "/commands/shipping/handover/scan", "Shipping Manage", "Memasukkan resi ke batch serah-terima"),
        ("POST", "/commands/shipping/handover/finalize", "Shipping Manage", "Finalisasi batch serah-terima"),
        ("POST", "/commands/pricing", "Owner/Admin", "Menyimpan konfigurasi HPP/marketplace"),
        ("POST/PATCH", "/commands/employees, /commands/employees/:id", "Owner", "Mengelola data kerja karyawan"),
        ("POST", "/commands/attendance", "Login", "Mencatat kehadiran sesuai lokasi"),
        ("POST", "/commands/loans, /commands/payrolls", "Owner", "Kasbon dan payroll"),
        ("POST", "/uploads/image", "Login", "Upload JPG/PNG/WEBP maksimal 5 MB"),
    ]
    add_table(doc, ["Method", "Path", "Akses", "Fungsi"], api_rows, [900, 3900, 1500, 3060], 8.0)

    doc.add_heading("4.1 Contoh request", level=2)
    examples = [
        ("Login", 'POST /api/login\n{ "email": "tester@menengsofficial.com", "password": "<password-tester>" }'),
        ("Stok masuk", 'POST /api/commands/receipts\n{ "locationId":"loc-id", "sourceType":"supplier", "supplierName":"Supplier UAT", "items":[{ "variantId":"var-id", "quantity":100, "unitCost":6000 }], "note":"UAT" }'),
        ("Penjualan", 'POST /api/commands/sales\n{ "locationId":"loc-id", "items":[{ "variantId":"var-id", "quantity":2, "price":10000 }], "channel":"offline", "payment":"Tunai", "note":"UAT" }'),
    ]
    for label, code in examples:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        font(p.add_run(label), size=10, bold=True, color=DARK_BLUE)
        table = doc.add_table(rows=1, cols=1)
        set_table_fixed(table, [9360])
        set_cell_shading(table.cell(0, 0), "F3F5F7")
        cp = table.cell(0, 0).paragraphs[0]
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.line_spacing = 1.0
        font(cp.add_run(code), size=8.5, name="Courier New", color=NAVY)
    doc.add_heading("4.2 Status dan error umum", level=2)
    add_table(doc, ["HTTP", "Makna", "Tindakan tester"], [
        ("200/201", "Berhasil / dibuat", "Verifikasi perubahan di UI, state, saldo, dan histori"),
        ("400", "Payload/validasi gagal", "Catat message; pastikan tidak ada perubahan parsial"),
        ("401", "Token tidak ada/kedaluwarsa", "Login ulang"),
        ("403", "Role/lokasi tidak berizin", "Ini expected untuk negative test"),
        ("404", "Endpoint/fitur tidak tersedia", "Catat URL dan langkah reproduksi"),
        ("409", "Konflik data/version", "Refresh data lalu ulangi; pastikan tidak double-submit"),
        ("413", "Payload terlalu besar", "Hubungi administrator untuk pengarsipan"),
        ("429", "Rate limit login", "Tunggu 15 menit; jangan brute-force"),
        ("503", "Database tidak dapat dihubungi", "Catat waktu, endpoint, dan screenshot"),
    ], [1100, 3100, 5160], 8.8)

    page_break(doc)
    doc.add_heading("5. Matriks Role dan Permission", level=1)
    add_p(doc, "Legenda: ✓ = diizinkan; S = hanya lokasi yang ditetapkan; R = baca saja; — = tidak tersedia. Permission backend tetap menjadi sumber kebenaran meskipun menu disembunyikan di UI.")
    perm_rows = [
        ("Produk lihat", "✓", "✓", "R/S", "R/S", "R/S", "—", "—"),
        ("Produk buat/ubah", "✓", "✓", "—", "—", "—", "—", "—"),
        ("Lokasi lihat", "✓", "✓", "S", "S", "S", "—", "—"),
        ("Lokasi buat/ubah", "✓", "✓", "—", "—", "—", "—", "—"),
        ("Kelola user/role", "✓", "—", "—", "—", "—", "—", "—"),
        ("Stok lihat", "✓", "✓", "S", "S", "S", "R", "—"),
        ("Stok masuk", "✓", "✓", "—", "S", "—", "—", "—"),
        ("Stock opname", "✓", "✓", "S", "S", "—", "—", "—"),
        ("Transfer kirim", "✓", "✓", "S", "S", "—", "—", "—"),
        ("Transfer terima", "✓", "✓", "S", "—", "—", "—", "—"),
        ("Penjualan buat", "✓", "✓", "S", "—", "S", "—", "—"),
        ("Void penjualan", "✓", "✓", "S", "—", "—", "—", "—"),
        ("Pengiriman", "✓", "✓", "S", "S", "—", "—", "—"),
        ("Laporan", "✓", "✓", "S", "S", "—", "R", "—"),
        ("Audit penuh", "✓", "—", "—", "—", "—", "—", "—"),
        ("Profil usaha", "✓", "—", "—", "—", "—", "—", "—"),
        ("Karyawan/kasbon/payroll", "✓", "—", "—", "—", "—", "—", "Absensi"),
    ]
    add_table(doc, ["Kapabilitas", "Owner", "Admin", "PIC", "Gudang", "Kasir", "Keu.", "Karyawan"], perm_rows, [2700, 900, 900, 900, 980, 900, 980, 1100], 7.6)
    doc.add_heading("5.1 Negative test wajib", level=2)
    negative = [
        ("Finance", "Coba POST penjualan/stok masuk", "403; tidak ada saldo/transaksi berubah"),
        ("Cashier", "Coba void, transfer, atau opname", "403 / aksi tidak tampil"),
        ("Warehouse", "Coba menerima transfer di outlet", "403 / lokasi tidak tersedia"),
        ("PIC Outlet A", "Akses data Outlet B", "Data tidak terlihat dan API menolak"),
        ("Admin", "Coba ubah profil usaha atau payroll", "403 / menu tidak tersedia"),
        ("Employee", "Buka menu selain Absensi/Bantuan", "Menu tidak tersedia"),
    ]
    add_table(doc, ["Role", "Uji", "Expected"], negative, [1600, 4000, 3760], 9)

    page_break(doc)
    doc.add_heading("6. Known Issues dan Limitation", level=1)
    limitations = [
        ("KL-01", "Auto-identifikasi ekspedisi tidak menjadi acuan utama", "Gunakan marketplace + resi manual; carrier dikonfirmasi saat serah-terima.", "Non-blocker"),
        ("KL-02", "Scan kamera bergantung HTTPS, izin browser, fokus, cahaya, dan kualitas label", "Selalu uji fallback input resi manual.", "Operasional"),
        ("KL-03", "Konversi satuan otomatis belum didukung", "Gunakan satuan konsisten antara produk, HPP, stok, dan transaksi.", "Diterima"),
        ("KL-04", "Pendaftaran mandiri production dinonaktifkan", "Akun dibuat Owner/administrator.", "Security"),
        ("KL-05", "Workspace terisolasi per organisasi; user satu organisasi berbagi data sesuai role", "Pastikan akun dibuat di workspace yang benar.", "By design"),
        ("KL-06", "Aplikasi membutuhkan koneksi internet; offline queue tidak tersedia", "Jangan memproses transaksi saat koneksi tidak stabil.", "Operasional"),
        ("KL-07", "Bundle JavaScript utama masih di atas rekomendasi 500 kB", "Pantau first-load pada perangkat/jaringan lambat; fungsi tidak terhambat.", "Non-blocker"),
        ("KL-08", "Payload state yang sangat besar dapat ditolak HTTP 413", "Arsipkan histori dan pertimbangkan pagination/server-side filtering saat volume meningkat.", "Skalabilitas"),
        ("KL-09", "Satu resep HPP dapat mengubah Harga Modal banyak varian", "Baca daftar dampak dan konfirmasi varian sebelum simpan.", "Kontrol user"),
        ("KL-10", "Penghapusan transaksi historis diganti pola pembatalan/void", "Masukkan alasan; verifikasi reversal pada saldo dan histori.", "Auditability"),
    ]
    add_table(doc, ["ID", "Limitation", "Mitigasi tester/operator", "Status"], limitations, [800, 3300, 4000, 1260], 8.3)
    doc.add_heading("6.1 Kapan pengujian harus dihentikan", level=2)
    add_bullet(doc, "Saldo menjadi negatif, berubah tanpa transaksi, atau reversal terjadi lebih dari sekali.")
    add_bullet(doc, "User dapat melihat data organisasi/lokasi lain atau melakukan aksi di luar role.")
    add_bullet(doc, "Data production nyata ikut berubah saat memakai workspace Tester.")
    add_bullet(doc, "Login/API mengembalikan error 5xx berulang pada alur Critical.")
    add_bullet(doc, "Terjadi kehilangan data setelah refresh, logout-login, atau deploy.")

    page_break(doc)
    doc.add_heading("7. Form Ringkasan Hasil dan Sign-off", level=1)
    add_table(doc, ["Metrik", "Hasil"], [
        ("Total skenario dieksekusi", ""),
        ("Lulus / Gagal / Blocked", ""),
        ("Defect Critical terbuka", ""),
        ("Defect High terbuka", ""),
        ("Tanggal mulai / selesai UAT", ""),
        ("Versi/deployment yang diuji", ""),
        ("Rekomendasi", "GO / CONDITIONAL GO / NO-GO"),
    ], [3300, 6060], 10)
    doc.add_heading("7.1 Daftar defect terbuka", level=2)
    add_table(doc, ["Issue ID", "AC ID", "Severity", "Ringkasan", "Owner", "Target", "Status"], [("", "", "", "", "", "", "") for _ in range(6)], [1100, 900, 1100, 2800, 1200, 1100, 1160], 8.5)
    doc.add_heading("7.2 Persetujuan", level=2)
    add_table(doc, ["Peran", "Nama", "Keputusan", "Tanggal", "Tanda tangan"], [
        ("Lead Tester", "", "", "", ""),
        ("Product Owner", "", "", "", ""),
        ("Developer", "", "", "", ""),
        ("Penerima Serah Terima", "", "", "", ""),
    ], [1700, 2000, 1900, 1600, 2160], 9)
    add_callout(doc, "Kriteria Go-Live", "Dokumen dapat ditandatangani GO hanya jika semua Critical lulus, tidak ada defect Critical/High terbuka, backup/restore tervalidasi, serta Product Owner menerima seluruh known limitation.", "ok")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
